from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt

output_path = "/Users/ibragimkuszeterov/Desktop/veb_main_univer/Django_cource_sem_4/onlineStore/Отчет_по_интеграциям_OAuth_Sentry.docx"

images = [
    {
        "path": "/Users/ibragimkuszeterov/Desktop/Снимок экрана 2026-05-20 в 18.30.50.png",
        "text": "На этом этапе я создал OAuth-клиент в Google Cloud и получил идентификатор клиента и секрет для подключения авторизации в проекте.",
        "caption": "Рисунок 45 - Создание OAuth client в Google Cloud"
    },
    {
        "path": "/Users/ibragimkuszeterov/Desktop/Снимок экрана 2026-05-20 в 18.27.47.png",
        "text": "Здесь я заполнил базовые поля брендинга OAuth-приложения, включая имя, контактный email и логотип проекта.",
        "caption": "Рисунок 46 - Настройка branding и логотипа OAuth-приложения"
    },
    {
        "path": "/Users/ibragimkuszeterov/Desktop/Снимок экрана 2026-05-20 в 18.28.54.png",
        "text": "На странице аудитории я проверил статус публикации и убедился, что тип пользователей выставлен как External для тестового доступа.",
        "caption": "Рисунок 47 - Проверка Audience и статуса публикации OAuth"
    },
    {
        "path": "/Users/ibragimkuszeterov/Desktop/Снимок экрана 2026-05-20 в 18.39.56.png",
        "text": "На этом шаге я добавил разрешенные JavaScript origins и redirect URI для локального запуска и корректного callback после входа через Google.",
        "caption": "Рисунок 48 - Конфигурация origins и redirect URI клиента OAuth"
    },
    {
        "path": "/Users/ibragimkuszeterov/Desktop/Снимок экрана 2026-05-20 в 18.50.15.png",
        "text": "Далее я открыл инструкцию Sentry для Django, чтобы сверить установку SDK, DSN и способ проверки отправки ошибок.",
        "caption": "Рисунок 49 - Экран настройки Sentry SDK для Django"
    },
    {
        "path": "/Users/ibragimkuszeterov/Desktop/download.png",
        "text": "Здесь я выбрал платформу Django в мастере Sentry и перешел к параметрам частоты уведомлений по ошибкам.",
        "caption": "Рисунок 50 - Выбор платформы Django и параметров алертов в Sentry"
    },
    {
        "path": "/Users/ibragimkuszeterov/Desktop/Снимок экрана 2026-05-20 в 19.05.09.png",
        "text": "После тестового запроса я получил список инцидентов в Sentry и увидел зафиксированную ошибку из debug endpoint.",
        "caption": "Рисунок 51 - Список зарегистрированных ошибок в Sentry"
    },
    {
        "path": "/Users/ibragimkuszeterov/Desktop/Снимок экрана 2026-05-20 в 19.05.43.png",
        "text": "В карточке инцидента я проверил детали события, окружение, пользователя и стек ошибки, чтобы подтвердить корректную интеграцию.",
        "caption": "Рисунок 52 - Детальная карточка события ошибки в Sentry"
    },
    {
        "path": "/Users/ibragimkuszeterov/Desktop/Снимок экрана 2026-05-20 в 19.05.59.png",
        "text": "На таймлайне события я просмотрел breadcrumbs и HTTP-данные запроса, что подтвердило полный трейс ошибки от API до Sentry.",
        "caption": "Рисунок 53 - Таймлайн и HTTP-контекст события в Sentry"
    },
]

doc = Document()

style = doc.styles["Normal"]
style.font.name = "Times New Roman"
style.font.size = Pt(14)

title = doc.add_paragraph("Отчет по настройке OAuth2 и Sentry")
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.runs[0].bold = True
title.runs[0].font.size = Pt(16)

doc.add_paragraph(
    "В этом разделе я показываю последовательность шагов, которые выполнил при настройке OAuth2 через Google и интеграции мониторинга ошибок через Sentry в своем проекте."
)

for item in images:
    p = doc.add_paragraph(item["text"])
    p.paragraph_format.space_after = Pt(8)

    img_p = doc.add_paragraph()
    img_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = img_p.add_run()
    run.add_picture(item["path"], width=Cm(16.5))

    cap = doc.add_paragraph(item["caption"])
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.runs[0].italic = True
    cap.paragraph_format.space_after = Pt(12)


doc.save(output_path)
print(output_path)
